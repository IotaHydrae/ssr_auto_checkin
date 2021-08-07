# coding:utf-8

from docx import Document
from docx.shared import Length, Pt, RGBColor
import os

# document = Document('examples/tkinter_emailv1.0.docx')
#
# for p in document.paragraphs:
#     print(p.text)
files_list = []
document =  Document()

def visitDir(path):
    list_all = os.listdir(path)

    for item in list_all:
        name = os.path.join(path, item)
        if not os.path.isfile(name):
            visitDir(name)
        else:
            if name != os.getcwd() +'\\'+ os.path.basename(__file__):
                # print(name)
                files_list.append(name)
    return files_list


def read_file(path):
    str = ''
    filename = path.split('\\')[-1]
    file_type = ['.c', '.py', '.java', '.S', '.dis', '.sh', '.cpp', '.h']
    if path.startswith('.'):
        return
    try:
        if os.path.splitext(path)[1] in file_type or filename == 'Makefile':
            with open(path, 'r', encoding='utf-8') as file:
                str = file.read()
                file.close()
    except:
        pass

    return str

def generate(filename, data):
    if len(data)>0:
        document.add_heading(filename)
        document.add_paragraph(data)




if __name__ == '__main__':

    files_list = visitDir(os.getcwd())
    for i in files_list:
        filename = i.split('\\')[-1]
        data = read_file(i)
        generate(filename, data)
    document.save('./{}.docx'.format(os.getcwd().split('\\')[-1]))
